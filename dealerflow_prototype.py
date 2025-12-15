import os
import json
import pandas as pd
from docx import Document
import PyPDF2
from gpt4all import GPT4All

# utility functions to read documents (doc, pdf, excel)
def read_docx(path):
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

def read_pdf(path):
    text = ""
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            if page.extract_text():
                text += page.extract_text() + "\n"
    return text

def read_excel(path):
    df = pd.read_excel(path, engine="openpyxl")
    return df.to_string(index=False)

# Chunking (for context control)
def chunk_text(text, max_chars=1200):
    chunks = []
    current = ""
    for line in text.split("\n"):
        if len(current) + len(line) < max_chars:
            current += line + "\n"
        else:
            chunks.append(current)
            current = line + "\n"
    if current.strip():
        chunks.append(current)
    return chunks

#summarization pipeline
def summarize_chunks(chunks, model):
    summaries = []
    for i, chunk in enumerate(chunks):
        print(f"Summarizing chunk {i+1}/{len(chunks)}")
        summary = model.generate(
            f"""
Summarize the following text focusing only on:
- Product capabilities
- Target customers
- Industry
- Business value
- Integrations

Text:
{chunk}
""",
            max_tokens=250
        )
        summaries.append(summary)
    return "\n".join(summaries)

#Final generation of companies
def generate_companies(company_type, region, context, model):
#     prompt = f"""
# You are an AI market analyst.

# Product context:
# {context}

# Task:
# Generate 10 companies that would most likely be {company_type}s for this product.

# Return ONLY valid JSON array.
# Each object must include:
# - name
# - website
# - locations
# - size
# - rationale

# Region preference: {region or "Any"}
# """
    prompt = f"""
You are an AI market analyst.

STRICT OUTPUT REQUIREMENTS:
- Output MUST be valid JSON
- Output MUST be a JSON array of EXACTLY 10 companies
- EACH company object MUST contain ALL fields below:
    - company_name (string)
    - website_url (string)
    - operating_locations (string or list)
    - estimated_size (one of: Small, Medium, Large)
    - rationale (string explaining why this company fits)

Rules:
- Use only real, publicly known companies
- Do NOT use LinkedIn or private data
- Base reasoning on product fit and scale

Product context:
{context}

Company type: {company_type}
Region preference: {region or "Any"}

Return ONLY JSON. No explanations.
"""
    return model.generate(prompt, max_tokens=900)

def save_output(output):
    try:
        data = json.loads(output)
        df = pd.DataFrame(data)
        df.to_csv("dealerflow_companies.csv", index=False)
        print("Saved: dealerflow_companies.csv")
    except Exception:
        print("JSON invalid: saving raw output")
        with open("dealerflow_companies_raw.txt", "w", encoding="utf-8") as f:
            f.write(output)
        print("Saved: dealerflow_companies_raw.txt")

if __name__ == "__main__":
    print("\nDealerFlow Cloud™ AI Prototype\n")

    company_type = input("Enter type (Customer / Partner): ").strip()
    region = input("Optional region (e.g., USA): ").strip() or None

    print("\nLoading local LLM (first run may take some time)")
    model = GPT4All("orca-mini-3b-gguf2-q4_0.gguf")

    docs_folder = "Fictitious_Company_AxelwaveTechnologies_DemoData"

    files = [
        "AxleWave_DealerFlowCloud_PRD.docx",
        "AxleWave_DealerFlowCloud_MRD.docx",
        "AxleWave_DealerFlowCloud_SAD.docx",
        "AxleWave_DealerFlowCloud_Technical_Docs.pdf",
        "AxleWave_DealerFlowCloud_User_Manual.docx",
        "AxleWave_DealerFlowCloud_Release_Notes.docx",
        "Axelwave_Technologies_Company_and_Product_Summary.docx",
        "AxleWave_DealerFlowCloud_Customer_Feedback_Log.xlsx"
    ]

    print("Reading documents")
    combined_text = ""

    for file in files:
        path = os.path.join(docs_folder, file)
        if file.endswith(".docx"):
            combined_text += read_docx(path)
        elif file.endswith(".pdf"):
            combined_text += read_pdf(path)
        elif file.endswith(".xlsx"):
            combined_text += read_excel(path)

    print("Chunking documents")
    chunks = chunk_text(combined_text)

    print("Creating condensed context")
    summary_context = summarize_chunks(chunks, model)

    print("Generating companies")
    output = generate_companies(company_type, region, summary_context, model)

    save_output(output)

    print("\nCompleted")



# # dealerflow_prototype_hf.py

# import os
# import json
# import pandas as pd
# from docx import Document
# import PyPDF2
# import requests

# # -------------------------------
# # Hugging Face API token
# # -------------------------------
# HF_API_TOKEN = os.getenv("HF_API_TOKEN")  # set in .env or environment
# API_URL = "https://huggingface.co/mosaicml/mpt-7b"

# headers = {"Authorization": f"Bearer {HF_API_TOKEN}"}

# # -------------------------------
# # Read docs
# # -------------------------------
# def read_docx(file_path):
#     from docx import Document
#     doc = Document(file_path)
#     return "\n".join([para.text for para in doc.paragraphs])

# def read_pdf(file_path):
#     text = ""
#     with open(file_path, "rb") as f:
#         reader = PyPDF2.PdfReader(f)
#         for page in reader.pages:
#             page_text = page.extract_text()
#             if page_text:
#                 text += page_text + "\n"
#     return text

# # -------------------------------
# # Read all documents
# # -------------------------------
# docs_folder = "Fictitious_Company_AxelwaveTechnologies_DemoData"

# doc_files = {
#     "PRD": "AxleWave_DealerFlowCloud_PRD.docx",
#     "MRD": "AxleWave_DealerFlowCloud_MRD.docx",
#     "SAD": "AxleWave_DealerFlowCloud_SAD.docx",
#     "TechDocs": "AxleWave_DealerFlowCloud_Technical_Docs.pdf",
#     "UserManual": "AxleWave_DealerFlowCloud_User_Manual.docx",
#     "ReleaseNotes": "AxleWave_DealerFlowCloud_Release_Notes.docx",
#     "Summary": "Axelwave_Technologies_Company_and_Product_Summary.docx",
#     "Feedback": "AxleWave_DealerFlowCloud_Customer_Feedback_Log.xlsx"
# }

# combined_text = ""

# # Docx and PDF
# for key, file in doc_files.items():
#     path = os.path.join(docs_folder, file)
#     if path.endswith(".docx"):
#         combined_text += read_docx(path) + "\n"
#     elif path.endswith(".pdf"):
#         combined_text += read_pdf(path) + "\n"
#     elif path.endswith(".xlsx"):
#         df = pd.read_excel(path, engine="openpyxl")
#         combined_text += df.to_string(index=False) + "\n"

# # -------------------------------
# # Hugging Face query
# # -------------------------------
# def query_hf(prompt):
#     payload = {"inputs": prompt}
#     response = requests.post(API_URL, headers=headers, json=payload)
#     try:
#         # some models return string, some JSON
#         output = response.json()
#         if isinstance(output, list) and "generated_text" in output[0]:
#             return output[0]["generated_text"]
#         elif isinstance(output, dict) and "error" in output:
#             return output["error"]
#         else:
#             return str(output)
#     except Exception as e:
#         return str(e)

# # -------------------------------
# # Generate companies
# # -------------------------------
# def generate_companies(company_type, region=None):
#     prompt = f"""
# You are an AI market analyst. You have the following information about a product:

# {combined_text}

# Task: Generate 10 companies most likely to be {company_type} for this product. 
# For each company, provide:
# - Name
# - Website
# - Locations where they operate
# - Size (Small, Medium, Large)
# - Rationale for why this company is relevant

# Only provide output in JSON array format.
# Region: {region if region else 'any region'}
# """
#     return query_hf(prompt)

# # -------------------------------
# # Save to CSV
# # -------------------------------
# def save_to_csv(companies_json, output_file="dealerflow_companies.csv"):
#     try:
#         companies_list = json.loads(companies_json)
#     except json.JSONDecodeError:
#         print("Error: output not valid JSON, saving raw text instead.")
#         with open("dealerflow_companies_raw.txt", "w") as f:
#             f.write(companies_json)
#         return
#     df = pd.DataFrame(companies_list)
#     df.to_csv(output_file, index=False)
#     print(f"✅ Output saved to {output_file}")

# # -------------------------------
# # Run Prototype
# # -------------------------------
# if __name__ == "__main__":
#     company_type = input("Enter type (Customer / Partner): ").strip()
#     region = input("Optional: Enter region (e.g., USA): ").strip() or None

#     print("Generating companies... this may take 30-60 seconds depending on model...")
#     companies_json = generate_companies(company_type, region)
#     save_to_csv(companies_json)



# # dealerflow_prototype.py

# import os
# import json
# import pandas as pd
# from docx import Document
# import PyPDF2
# from dotenv import load_dotenv
# import openai

# # -------------------------------
# # Step 2a: Load OpenAI API Key
# # -------------------------------
# load_dotenv()
# OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
# openai.api_key = OPENAI_API_KEY

# # -------------------------------
# # Step 2b: Functions to read docs
# # -------------------------------
# def read_docx(file_path):
#     doc = Document(file_path)
#     return "\n".join([para.text for para in doc.paragraphs])

# def read_pdf(file_path):
#     text = ""
#     with open(file_path, "rb") as f:
#         reader = PyPDF2.PdfReader(f)
#         for page in reader.pages:
#             page_text = page.extract_text()
#             if page_text:
#                 text += page_text + "\n"
#     return text

# def read_excel(file_path):
#     df = pd.read_excel(file_path)
#     return df.to_string(index=False)

# # -------------------------------
# # Step 2c: Load all documents
# # -------------------------------
# docs_folder = "Fictitious_Company_AxelwaveTechnologies_DemoData"

# prd_text = read_docx(os.path.join(docs_folder, "AxleWave_DealerFlowCloud_PRD.docx"))
# mrd_text = read_docx(os.path.join(docs_folder, "AxleWave_DealerFlowCloud_MRD.docx"))
# sad_text = read_docx(os.path.join(docs_folder, "AxleWave_DealerFlowCloud_SAD.docx"))
# tech_text = read_pdf(os.path.join(docs_folder, "AxleWave_DealerFlowCloud_Technical_Docs.pdf"))
# user_manual_text = read_docx(os.path.join(docs_folder, "AxleWave_DealerFlowCloud_User_Manual.docx"))
# release_notes_text = read_docx(os.path.join(docs_folder, "AxleWave_DealerFlowCloud_Release_Notes.docx"))
# summary_text = read_docx(os.path.join(docs_folder, "Axelwave_Technologies_Company_and_Product_Summary.docx"))

# feedback_text = read_excel(os.path.join(docs_folder, "AxleWave_DealerFlowCloud_Customer_Feedback_Log.xlsx"))

# # Combine all text into one large string for GPT
# combined_docs_text = "\n".join([
#     summary_text, prd_text, mrd_text, sad_text, tech_text, user_manual_text, release_notes_text, feedback_text
# ])

# # -------------------------------
# # Step 2d: GPT function to generate companies
# # -------------------------------
# def generate_companies(company_type, region=None):
#     prompt = f"""
# You are an AI market analyst. You have the following information about a product:

# {combined_docs_text}

# Task: Generate 10 companies most likely to be {company_type} for this product. 
# For each company, provide:
# - Name
# - Website
# - Locations where they operate
# - Size (Small, Medium, Large)
# - Rationale for why this company is relevant

# Only provide output in JSON array format.
# Region: {region if region else 'any region'}
# """
#     response = openai.chat.completions.create(
#         model="gpt-5-mini",
#         messages=[{"role": "user", "content": prompt}],
#         temperature=0.7
#     )
#     return response.choices[0].message.content


# # -------------------------------
# # Step 2e: Save output to CSV
# # -------------------------------
# def save_companies_to_csv(companies_json, output_file="dealerflow_companies.csv"):
#     try:
#         companies_list = json.loads(companies_json)
#     except json.JSONDecodeError:
#         print("Error: GPT output was not valid JSON. Printing raw output:")
#         print(companies_json)
#         return

#     df = pd.DataFrame(companies_list)
#     df.to_csv(output_file, index=False)
#     print(f"✅ Output saved to {output_file}")

# # -------------------------------
# # Step 3: Run Prototype
# # -------------------------------
# if __name__ == "__main__":
#     print("DealerFlow Cloud™ AI Prototype")
#     company_type = input("Enter type (Customer / Partner): ").strip()
#     region = input("Optional: Enter region (e.g., USA): ").strip() or None

#     print("Generating companies... this may take 30-60 seconds depending on API...")
#     companies_json = generate_companies(company_type, region)

#     save_companies_to_csv(companies_json)




# from dotenv import load_dotenv
# import os
# load_dotenv()
# OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# company_type = input("Enter type (Customer / Partner): ")
# region = input("Optional: enter region (e.g., USA): ")
# size = input("Optional: enter size (Small/Medium/Large): ")

# import openai

# def generate_companies(company_type, region=None):
#     prompt = f"""
#     You are an AI market analyst. I have a fictitious company DealerFlow Cloud™, a cloud-native platform
#     for automotive dealers, OEMs, and partners. Your task is to generate 10 {company_type} companies 
#     most likely to adopt or partner with this product. Include for each company:
#     - Name
#     - Website
#     - Locations
#     - Size (Small/Medium/Large)
#     - Rationale

#     Only provide output in JSON array format.
#     Region: {region if region else 'any'}
#     """
#     response = openai.ChatCompletion.create(
#         model="gpt-5-mini",
#         messages=[{"role": "user", "content": prompt}],
#         temperature=0.7
#     )
#     companies_json = response['choices'][0]['message']['content']
#     return companies_json

# companies = generate_companies("Customer", "USA")
# print(companies)

# import requests
# from bs4 import BeautifulSoup

# def get_company_website(name):
#     search_url = f"https://www.google.com/search?q={name}+official+site"
#     # For prototype, we can simply return placeholder or instruct GPT to fill it
#     return f"https://www.{name.replace(' ', '').lower()}.com"

# import pandas as pd
# import json

# # Assuming `companies` is JSON string
# companies_list = json.loads(companies)
# df = pd.DataFrame(companies_list)
# df.to_csv("dealerflow_companies.csv", index=False)
# print("Output saved to dealerflow_companies.csv")

# def verify_companies(companies_json):
#     prompt = f"""
#     You are a market verification AI. Check the following companies if they match the 
#     target segment for DealerFlow Cloud™ (dealers, OEMs, automotive partners). Return only companies 
#     that are highly relevant with a score 1-5:
#     {companies_json}
#     """
#     response = openai.ChatCompletion.create(
#         model="gpt-5-mini",
#         messages=[{"role": "user", "content": prompt}],
#         temperature=0
#     )
#     return response['choices'][0]['message']['content']

# verified_companies = verify_companies(companies)
# print(verified_companies)

